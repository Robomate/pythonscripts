#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
ST-Twincat2-AutoMapping
Purpose: Load EXP Variablen di and do from exp file
Version: 01/2020 Roboball
"""
import sys
import os
import docx
from copy import deepcopy
import win32com.client
 



# define global
project_name = "SKE_P_PLC03"

# export doc
# data_file_export = "Test_V1.docx"
data_file_export = "IBN Checkliste.docx"

# set template name
sheet_name_01 = "Template_Checkliste__Empty.docx"

# paths
path_01 = "C:\\Users\" + project_name + "\\Create_Checkliste\\_templates\\"
path_02 = "C:\\Users\" + project_name + "\\Create_Checkliste\\0_Checklisten\\"
path_03 = "C:\\Users\" + project_name + "\\Create_Project_Mappings\\"

# add path to import helper files
sys.path.append(path_03)
import _helper_globals_02 as hg
import _helper_functions_03 as hf

# global vars

# for table: Deckblatt 0
kunde =  "SKEMOR"
projectnr = "P1505"
auftragsnr = "520229500"
datum = "09.03.2021"
plcler = "XXXX"
plcnr = "03"

convert_cm_inch = 0.393701 # (1cm = 0.393701inch)
convert_cm_point = 28.3465 # (1cm = 28.3465points)

list_table_names = [
    "Deckblatt",
    "Vorbereitung",
    "Allgemein 1",
    "Rolladapter 500_KF",
    "Rolladapter 500_F",
    "Rolladapter 500_STF",
    "200G_KF",
    "Rolladapter 500_Funktionen",
    "200G_Funktionen",
    "Allgemein 2",
    "WCS-Test",
    "Allgemeine Informationen",
    "Final",
]


# define doxc functions
def print_tables(doc):
    ''' print all tables of a doc file '''
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    print(paragraph.text)
                print_tables(cell)

def read_in_docx_table_by_col(
                            doc, # full doc
                            table_idx, # specific table
                            print_opt, # optional print
                            ):
    ''' 
    read in single table by column
    return: list of columns
    dependencies: docx
    '''
    num_rows = len(doc.tables[table_idx].column_cells(0))
    num_cols = len(doc.tables[table_idx].row_cells(0))
    list_single_table = []
    for idx_1 in range(num_cols):
        list_single_col = []
        for obj in doc.tables[table_idx].column_cells(idx_1):
            list_single_col.append(obj.text)
        list_single_table.append(list_single_col)
    if print_opt: # optional : print full table
        print("Table idx: " + str(table_idx) + ", size(rows,cols): (" +\
             str(num_rows) + "," + str(num_cols) + ")")
        print(*list_single_table, sep = "\n")
        print()
    return list_single_table

def read_in_docx_table_by_row(
                            doc, # full doc
                            table_idx, # specific table
                            print_opt, # optional print
                            ):
    ''' 
    read in single table by row
    return: list of row
    dependencies: docx
    '''
    num_rows = len(doc.tables[table_idx].column_cells(0))
    num_cols = len(doc.tables[table_idx].row_cells(0))
    list_single_table = []
    for idx_1 in range(num_rows):
        list_single_row = []
        for obj in doc.tables[table_idx].row_cells(idx_1):
            list_single_row.append(obj.text)
        list_single_table.append(list_single_row)
    if print_opt: # optional : print full table
        print("Table idx: " + str(table_idx) + ", size(rows,cols): (" +\
             str(num_rows) + "," + str(num_cols) + ")")
        print(*list_single_table, sep = "\n")
        print()
    return list_single_table

def fill_docx_table_by_col(
                            doc, # full doc
                            table_idx, # specific table
                            list_of_cols, # string data for cols
                            ):
    ''' 
    fill table with strings by columns
    return: doc with filled table object
    dependencies: docx
    '''
    rows = len(list_of_cols[0])
    cols = len(list_of_cols)
    for idx_01 in range(cols):
        col_temp = doc.tables[0].columns[idx_01].cells 
        for idx_02 in range(rows):
            col_temp[idx_02].text = list_of_cols[idx_01][6]
    return doc


def read_in_wincom32_word_table_by_row(
                                        doc, # full doc
                                        table_idx, # specific table
                                        print_opt, # optional print
                                        ):
    ''' 
    read in single table by row
    return: list of row
    dependencies: win32com.client
    '''
    num_rows = doc.Tables(table_idx).Rows.Count
    num_cols = doc.Tables(table_idx).Columns.Count
    list_single_table = []
    for row in range(1, num_rows + 1):
        list_single_row = []
        for col in range(1, num_cols + 1):
            try:
                data_str = doc.Tables(table_idx).Cell(Row = row,Column = col).Range.Text.split('\r\x07')[0]
            except:
                data_str = "BugFill" # "doc.Tables(table_idx).Cell(Row = row,Column = col).Range.Text"
            list_single_row.append(data_str)
        list_single_table.append(list_single_row)
    if print_opt: # optional : print full table
        print("Table idx: " + str(table_idx) + ", size(rows,cols): (" +\
             str(num_rows) + "," + str(num_cols) + ")")
        print(*list_single_table, sep = "\n")
        print()
    return list_single_table

def read_in_wincom32_word_table_by_col(
                                        doc, # full doc
                                        table_idx, # specific table
                                        print_opt, # optional print
                                        ):
    ''' 
    read in single table by row
    return: list of row
    dependencies: win32com.client
    '''
    num_rows = doc.Tables(table_idx).Rows.Count
    num_cols = doc.Tables(table_idx).Columns.Count
    list_single_table = []
    for col in range(1, num_cols + 1):
        list_single_col = []
        for row in range(1, num_rows + 1):
            try:
                data_str = doc.Tables(table_idx).Cell(Row = row,Column = col).Range.Text.split('\r\x07')[0]
            except:
                data_str = "BugFill" # "doc.Tables(table_idx).Cell(Row = row,Column = col).Range.Text"
            list_single_col.append(data_str)
        list_single_table.append(list_single_col)
    if print_opt: # optional : print full table
        print("Table idx: " + str(table_idx) + ", size(rows,cols): (" +\
             str(num_rows) + "," + str(num_cols) + ")")
        print(*list_single_table, sep = "\n")
        print()
    return list_single_table

def copy_content_btwn_word_docs_wincom32_(
                                        doc_ex, # word doc to copy content from
                                        doc_im, # word doc to copy content to
                                        ):
    ''' 
    copy entire content from word doc into other word doc
    return: doc_im
    dependencies: win32com.client
    '''
    doc_ex.Activate() # activate word doc to copy content from
    doc_ex.Application.Selection.WholeStory() # choose all
    doc_ex.Application.Selection.Expand(win32com.client.constants.wdParagraph) # expands the selection to current paragraph
    doc_ex.Application.Selection.Copy() # copy the selection
    doc_im.Activate() # activate word doc to copy content from
    doc_im.Range().Collapse(win32com.client.constants.wdCollapseEnd) # move to the end of the document
    doc_ex.Application.Selection.PasteAndFormat(win32com.client.constants.wdPasteDefault) # paste all incl. formatting

    return doc_im


if __name__ == '__main__':
    # https://fortes-arthur.medium.com/handling-with-doc-extension-with-python-b6491792311e
    # https://www.programmersought.com/article/8923249125/
    # https://stackoverflow.com/questions/63503373/duplicating-a-table-in-word-document-opened-using-win32-in-python

    ##############################################################
    # read-in by win32com.client
    wordapp = win32com.client.gencache.EnsureDispatch("Word.Application")
    # wordapp.Visible = True
    wordapp.Visible = False

    # open template
    doc_01 = wordapp.Documents.Open(path_01+sheet_name_01)

    # get number of sheets
    doc_01.Repaginate()
    num_of_pages = doc_01.ComputeStatistics(2)
    print("Template num of pages:", num_of_pages)

    # #############################################################
    # # optional: export all as empty templates:
    # for idx_01 in range(len(list_table_names)):
    #     doc_01.SaveAs(path_01 + "Template_Checkliste_" + str(idx_01) + "_" +list_table_names[idx_01] + ".docx")

    # ##############################################################
    # # do data storage from imported template
    # ##############################################################
    # num_tables = doc_01.Tables.Count # count tables in doc
    # print("\nRead in all Tables, Number: " + str(num_tables) + "\n")
    # list_doc_all_tables = []
    # for idx_01 in range(1,num_tables + 1,1):
    #     # use for print out only
    #     list_single_table_00 = read_in_wincom32_word_table_by_row(
    #                                                             doc_01, # full doc
    #                                                             idx_01, # start with 1, specific table
    #                                                             True, # optional print
    #                                                             )
        
    #     # use for import of templates
    #     list_single_table_01 = read_in_wincom32_word_table_by_col(
    #                                                             doc_01, # full doc
    #                                                             idx_01, # start with 1, specific table
    #                                                             False, # optional print
    #                                                             )
    #     list_doc_all_tables.append(list_single_table_01)



    ##############################################################
    # create export from template
    ##############################################################

    # remove earlier version of file
    if os.path.exists(path_02 + data_file_export):
        os.remove(path_02 + data_file_export)

    doc_02 = wordapp.Documents.Add() # Create new Document Object

    # add background template
    doc_02.Application.Selection.Range.InsertFile(path_01+sheet_name_01) # add background template

    ##############################################################
    # page setup
    doc_02.Activate() # Activate the other document 
    doc_02.PageSetup.Orientation = 1 # 1: horizontal "Landscape", 0: vertical (normal)
    # enter in cm: 
    doc_02.PageSetup.TopMargin = 2.0 * convert_cm_point
    doc_02.PageSetup.BottomMargin = 0.5 * convert_cm_point
    doc_02.PageSetup.LeftMargin = 1.5 * convert_cm_point
    doc_02.PageSetup.RightMargin = 2.8 * convert_cm_point

    # footer/header
    doc_02.PageSetup.DifferentFirstPageHeaderFooter = True
    doc_02.PageSetup.HeaderDistance = 0.7 * convert_cm_point 
    # doc_02.PageSetup.FooterDistance = 0.7 * convert_cm_point
    # tabstops
    doc_02.Content.Paragraphs.TabStops.Add(14.25 * convert_cm_point)
    doc_02.Content.Paragraphs.TabStops.Add(18.25 * convert_cm_point)

    # apply styles
    doc_02.Application.Selection = wordapp.ActiveDocument.Styles("Kein Leerraum")
    doc_02.Application.Selection.TypeParagraph() # THIS IS WHAT I ADDED, NO NEED TO READJUST STYLE

    ##############################################################
    # add templates
    ##############################################################
    print("\nTemplate Overview:")
    for idx_01 in range(len(list_table_names)):
        print(list_table_names[idx_01], idx_01)
    
    print("\nStart adding Templates:")
    ##############################################################
    # create intro
    for idx_01 in range(0,3,1):
        print(list_table_names[idx_01], idx_01)
        # add content template
        doc_03 = wordapp.Documents.Open(path_01 + "Template_Checkliste_" + str(idx_01) + "_" + list_table_names[idx_01] + ".docx") # open template
        doc_02 = copy_content_btwn_word_docs_wincom32_(
                                                    doc_03, # word doc to copy content from
                                                    doc_02, # word doc to copy content to
                                                    )
        doc_03.Close() # close template

        if idx_01 == 0: # fill table: Deckblatt 0
            doc_02.Tables(1).Cell(Row = 1,Column = 2).Range.Text = kunde # "SKEMOR"
            doc_02.Tables(1).Cell(Row = 2,Column = 2).Range.Text = projectnr # "P1505"
            doc_02.Tables(1).Cell(Row = 3,Column = 2).Range.Text = auftragsnr # "520229500"
            doc_02.Tables(1).Cell(Row = 4,Column = 2).Range.Text = datum # "06.03.2021"
            doc_02.Tables(1).Cell(Row = 5,Column = 2).Range.Text = plcler # "XXXX"
            doc_02.Tables(1).Cell(Row = 7,Column = 2).Range.Text = plcnr # "03"
    
    ##############################################################
    # create: 500 Rolladapter

    # create KFs:
    set_idx = 3
    for idx_01 in range(len(hg.UnitID_list_KF)):
        print(list_table_names[set_idx], set_idx, " KF:",idx_01+1)
        num_tables = doc_02.Tables.Count # count tables in doc
        # add content template
        doc_03 = wordapp.Documents.Open(path_01 + "Template_Checkliste_" + str(set_idx) + "_" + list_table_names[set_idx] + ".docx") # open template
        doc_02 = copy_content_btwn_word_docs_wincom32_(
                                                    doc_03, # word doc to copy content from
                                                    doc_02, # word doc to copy content to
                                                    )
        doc_03.Close() # close template
        # fill table: KFs
        doc_02.Tables(num_tables+1).Cell(Row = 1,Column = 4).Range.Text = hg.UnitID_list_KF[idx_01]
        doc_02.Tables(num_tables+1).Cell(Row = 2,Column = 4).Range.Text = hg.Bezeichnung_list_KF[idx_01] + "   Nr.: " + str(idx_01+1)
 
    # create PFs:
    set_idx = 4
    for idx_01 in range(len(hg.UnitID_list_PF)):
        print(list_table_names[set_idx], set_idx, " PF:",idx_01+1)
        num_tables = doc_02.Tables.Count # count tables in doc
        # add content template
        doc_03 = wordapp.Documents.Open(path_01 + "Template_Checkliste_" + str(set_idx) + "_" + list_table_names[set_idx] + ".docx") # open template
        doc_02 = copy_content_btwn_word_docs_wincom32_(
                                                    doc_03, # word doc to copy content from
                                                    doc_02, # word doc to copy content to
                                                    )
        doc_03.Close() # close template
        # fill table: PFs
        doc_02.Tables(num_tables+1).Cell(Row = 1,Column = 4).Range.Text = hg.UnitID_list_PF[idx_01]
        doc_02.Tables(num_tables+1).Cell(Row = 2,Column = 4).Range.Text = hg.Bezeichnung_list_PF[idx_01] + "   Nr.: " + str(idx_01+1)
    
    # create SFs:
    set_idx = 5
    for idx_01 in range(len(hg.UnitID_list_SF)):
        print(list_table_names[set_idx], set_idx, " SF:",idx_01+1)
        num_tables = doc_02.Tables.Count # count tables in doc
        # add content template
        doc_03 = wordapp.Documents.Open(path_01 + "Template_Checkliste_" + str(set_idx) + "_" + list_table_names[set_idx] + ".docx") # open template
        doc_02 = copy_content_btwn_word_docs_wincom32_(
                                                    doc_03, # word doc to copy content from
                                                    doc_02, # word doc to copy content to
                                                    )
        doc_03.Close() # close template
        # fill table: KFs
        doc_02.Tables(num_tables+1).Cell(Row = 1,Column = 4).Range.Text = hg.UnitID_list_SF[idx_01]
        doc_02.Tables(num_tables+1).Cell(Row = 2,Column = 4).Range.Text = "Steilfoerderer" + "   Nr.: " + str(idx_01+1)


    # functions: 500 Rolladapter
    set_idx = 7
    print(list_table_names[set_idx], set_idx)
    # add content template
    doc_03 = wordapp.Documents.Open(path_01 + "Template_Checkliste_" + str(set_idx) + "_" + list_table_names[set_idx] + ".docx") # open template
    doc_02 = copy_content_btwn_word_docs_wincom32_(
                                                doc_03, # word doc to copy content from
                                                doc_02, # word doc to copy content to
                                                )
    doc_03.Close() # close template
    
    ##############################################################
    # create outro
    # print("\nStart adding Templates:")
    for idx_01 in range(9,13,1):
        print(list_table_names[idx_01], idx_01)
        # add content template
        doc_03 = wordapp.Documents.Open(path_01 + "Template_Checkliste_" + str(idx_01) + "_" + list_table_names[idx_01] + ".docx") # open template
        doc_02 = copy_content_btwn_word_docs_wincom32_(
                                                    doc_03, # word doc to copy content from
                                                    doc_02, # word doc to copy content to
                                                    )
        doc_03.Close() # close template



    
    # add page-breaks or section breaks (after tables)
    doc_02.Activate() # Activate the other document 
    num_tables = doc_02.Tables.Count # count tables in doc
    print("\nApply Styles to Doc: Section Breaks")
    print("Total Number of Tables in Export Doc:", num_tables)
    for idx_01 in range(1,doc_02.Tables.Count+1,1): 
        set_break_type = win32com.client.constants.wdSectionBreakNextPage # insert section break
        doc_02.Range(doc_02.Tables(idx_01).Range.End,doc_02.Tables(idx_01).Range.End).InsertBreak(set_break_type) # insert break
    
    
    ##############################################################
    # save export, close and quit 
    print("\nSave, Export and Quit App:")
    doc_02.SaveAs(path_02 + data_file_export)
    doc_02.Close()
    doc_01.Close()
    wordapp.Quit()
    print("** doc successfully exported. **")









